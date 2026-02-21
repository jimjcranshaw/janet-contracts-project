"""
Generate a detailed Word report with LLM-generated rationales.
For each charity, analyses the top matches AND notable rejections
using the LLMMatchAnalyzer to produce reasoned explanations.
"""
import sys
import os
import logging

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import SessionLocal, settings
from app.models import ServiceProfile, Notice, NoticeMatch
from app.services.matching.llm_match_analyzer import LLMMatchAnalyzer

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Verdict colours for visual clarity
VERDICT_COLOURS = {
    "STRONG_MATCH": RGBColor(0x1B, 0x7A, 0x2B),   # Green
    "GOOD_MATCH":   RGBColor(0x2E, 0x7D, 0x32),    # Light green
    "PARTIAL_MATCH": RGBColor(0xF5, 0x7F, 0x17),    # Amber
    "NOT_SUITABLE":  RGBColor(0xC6, 0x28, 0x28),    # Red
}

TOP_MATCHES = 5     # Top N matches per charity to analyse in detail
TOP_REJECTIONS = 3  # Top N rejections per charity to explain


def add_styled_paragraph(doc, text, bold=False, italic=False, color=None, size=None):
    """Helper to add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def generate_report():
    db = SessionLocal()
    analyzer = LLMMatchAnalyzer(db)
    charities = db.query(ServiceProfile).all()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ─── Title Page ───
    doc.add_paragraph()  # spacing
    title = doc.add_heading('Grants AI — Matching Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph('AI-Driven Procurement Matching for the UK Charity Sector')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].italic = True
    
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Date: 17 February 2026\n').bold = True
    meta.add_run(f'Charities Analysed: {len(charities)}\n')
    meta.add_run('Data Source: Find a Tender Service (FTS) — Live Procurement Notices\n')
    meta.add_run('Analysis Engine: LLM-Reasoned Matching (GPT-4o-mini)')
    doc.add_page_break()

    # ─── Methodology ───
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        "This report combines two layers of analysis for each charity–tender pair:"
    )
    doc.add_paragraph(
        "Layer 1 — Automated Pre-Screening: A mechanical filter using vector embeddings "
        "(semantic similarity), CPV code overlap (domain relevance), and geographic proximity "
        "to produce a ranked shortlist. This rapidly narrows ~300 tenders to a manageable set.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Layer 2 — LLM-Reasoned Analysis: For shortlisted matches and borderline cases, "
        "a Large Language Model reviews the charity's full evidence base (mission, activities, "
        "beneficiary groups, income, regions) against the tender's requirements, and produces "
        "a structured verdict with a written rationale explaining WHY the match is or isn't suitable.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Verdicts are categorised as: STRONG_MATCH (high confidence, bid recommended), "
        "GOOD_MATCH (solid fit, worth pursuing), PARTIAL_MATCH (some alignment but gaps exist), "
        "or NOT_SUITABLE (fundamental misalignment or viability concern)."
    )
    doc.add_page_break()

    # ─── Per-Charity Sections ───
    for idx, charity in enumerate(charities):
        income_str = f"£{charity.latest_income:,.0f}" if charity.latest_income else "Not reported"
        logger.info(f"[{idx+1}/{len(charities)}] Analysing {charity.name}...")
        
        doc.add_heading(f'{charity.name}', level=1)
        
        # Profile summary
        doc.add_heading('Organisation Profile', level=2)
        profile_table = doc.add_table(rows=0, cols=2)
        profile_table.style = 'Table Grid'
        
        profile_fields = [
            ("Annual Income", income_str),
            ("Mission", (charity.mission or "Not available")[:300]),
            ("Activities", (charity.programs_services or "Not available")[:300]),
            ("Beneficiaries", ", ".join(charity.beneficiary_groups) if charity.beneficiary_groups else "General public"),
            ("Regions", ", ".join(charity.service_regions) if charity.service_regions else "Not specified"),
            ("CPV Codes", ", ".join(charity.inferred_cpv_codes) if charity.inferred_cpv_codes else "None inferred"),
        ]
        for label, value in profile_fields:
            row = profile_table.add_row()
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
            row.cells[1].text = value

        # ─── TOP MATCHES ───
        doc.add_heading('Recommended Tenders', level=2)
        
        top_matches = db.query(NoticeMatch, Notice).join(
            Notice, NoticeMatch.notice_id == Notice.ocid
        ).filter(
            NoticeMatch.org_id == charity.org_id
        ).order_by(NoticeMatch.score.desc()).limit(TOP_MATCHES).all()

        if not top_matches:
            doc.add_paragraph("No matches found in the current tender pool for this organisation.")
        else:
            for rank, (m, n) in enumerate(top_matches, 1):
                score = float(m.score) if m.score else 0.0
                
                # Call LLM for reasoned analysis
                mechanical = {
                    "semantic": f"{float(m.score_semantic or 0):.2f}",
                    "domain": f"{float(m.score_domain or 0):.2f}",
                    "geo": f"{float(m.score_geo or 0):.2f}",
                    "viability": m.viability_warning or "None",
                }
                analysis = analyzer.analyze_match(charity.org_id, n.ocid, mechanical)
                
                verdict = analysis.get("verdict", "PARTIAL_MATCH")
                verdict_color = VERDICT_COLOURS.get(verdict, RGBColor(0, 0, 0))
                
                # Tender heading
                h = doc.add_heading(level=3)
                h.add_run(f"#{rank}  ").font.size = Pt(12)
                run = h.add_run(n.title[:120])
                run.font.size = Pt(12)
                
                # Verdict badge
                p = doc.add_paragraph()
                vrun = p.add_run(f"  {verdict.replace('_', ' ')}  ")
                vrun.bold = True
                vrun.font.color.rgb = verdict_color
                vrun.font.size = Pt(12)
                conf = analysis.get("confidence", 0)
                p.add_run(f"   Confidence: {conf*100:.0f}%")
                
                # Scores table
                scores_table = doc.add_table(rows=2, cols=4)
                scores_table.style = 'Light Shading Accent 1'
                scores_table.rows[0].cells[0].text = "Semantic"
                scores_table.rows[0].cells[1].text = "Domain (CPV)"
                scores_table.rows[0].cells[2].text = "Geographic"
                scores_table.rows[0].cells[3].text = "Overall"
                scores_table.rows[1].cells[0].text = mechanical["semantic"]
                scores_table.rows[1].cells[1].text = mechanical["domain"]
                scores_table.rows[1].cells[2].text = mechanical["geo"]
                scores_table.rows[1].cells[3].text = f"{score:.2f}"
                
                # Rationale — the key addition
                doc.add_paragraph()
                rationale_p = doc.add_paragraph()
                rationale_p.add_run("Rationale: ").bold = True
                rationale_p.add_run(analysis.get("rationale", "Analysis not available."))
                
                # Strengths
                strengths = analysis.get("strengths", [])
                if strengths:
                    doc.add_paragraph()
                    sp = doc.add_paragraph()
                    sp.add_run("Strengths: ").bold = True
                    for s in strengths:
                        doc.add_paragraph(s, style='List Bullet')
                
                # Risks
                risks = analysis.get("risks", [])
                if risks:
                    rp = doc.add_paragraph()
                    rp.add_run("Risks / Gaps: ").bold = True
                    for r in risks:
                        doc.add_paragraph(r, style='List Bullet')
                
                # Recommendation
                rec = analysis.get("recommendation", "")
                if rec:
                    recp = doc.add_paragraph()
                    recp.add_run("Recommendation: ").bold = True
                    recp.add_run(rec).italic = True
                
                # Renewal Radar / Strategic Coach
                radar = m.risk_flags.get("renewal_radar") if m.risk_flags else None
                if radar:
                    doc.add_paragraph()
                    radar_p = doc.add_paragraph()
                    radar_p.add_run("Strategic Renewal Radar").bold = True
                    radar_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    summary = radar.get("radar_summary", "")
                    if summary:
                        doc.add_paragraph(summary, style='List Bullet')
                    
                    # Add strategic timeline advice
                    incumbent = radar.get("incumbent")
                    cycle = radar.get("estimated_cycle_years") or 3
                    
                    if incumbent:
                        advice = f"Strategy: Engage with {n.buyer.canonical_name if n.buyer else 'buyer'} to understand their relationship with {incumbent}. Build evidence on where you outperform this incumbent."
                        doc.add_paragraph(advice, style='List Bullet')
                
                # Viability warning
                if m.viability_warning:
                    wp = doc.add_paragraph()
                    wr = wp.add_run(f"⚠ {m.viability_warning}")
                    wr.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
                    wr.bold = True
                
                doc.add_paragraph()  # spacing

        # ─── REJECTIONS / NOT SUITABLE ───
        doc.add_heading('Notable Rejections & Risk Warnings', level=2)
        doc.add_paragraph(
            "These tenders had some degree of relevance but were flagged as unsuitable "
            "due to viability concerns, geographic mismatch, or fundamental misalignment."
        )
        
        # Get tenders with viability warnings or low scores despite some semantic relevance
        rejections = db.query(NoticeMatch, Notice).join(
            Notice, NoticeMatch.notice_id == Notice.ocid
        ).filter(
            NoticeMatch.org_id == charity.org_id,
            NoticeMatch.viability_warning.isnot(None)
        ).order_by(NoticeMatch.score_semantic.desc()).limit(TOP_REJECTIONS).all()

        if not rejections:
            # Fall back to lowest-scoring matches from the set
            rejections = db.query(NoticeMatch, Notice).join(
                Notice, NoticeMatch.notice_id == Notice.ocid
            ).filter(
                NoticeMatch.org_id == charity.org_id
            ).order_by(NoticeMatch.score.asc()).limit(TOP_REJECTIONS).all()

        for m, n in rejections:
            score = float(m.score) if m.score else 0.0
            
            # LLM analysis for rejection reasoning
            mechanical = {
                "semantic": f"{float(m.score_semantic or 0):.2f}",
                "domain": f"{float(m.score_domain or 0):.2f}",
                "geo": f"{float(m.score_geo or 0):.2f}",
                "viability": m.viability_warning or "None",
            }
            analysis = analyzer.analyze_match(charity.org_id, n.ocid, mechanical)
            
            p = doc.add_paragraph()
            p.add_run(f"✗ {n.title[:100]}").bold = True
            p.add_run(f"  (Score: {score:.2f})")
            
            # Rejection rationale
            rp = doc.add_paragraph()
            rp.add_run("Why Not Suitable: ").bold = True
            rp.add_run(analysis.get("rationale", m.viability_warning or "Low overall alignment."))
            
            risks = analysis.get("risks", [])
            if risks:
                for r in risks:
                    doc.add_paragraph(f"• {r}", style='List Bullet')
            
            doc.add_paragraph()  # spacing

        doc.add_page_break()

    # ─── Conclusion ───
    doc.add_heading('Appendix: System Configuration', level=1)
    doc.add_paragraph(f"Total Tenders in Pool: ~300 (FTS live data)")
    doc.add_paragraph(f"Total Charities Analysed: {len(charities)}")
    doc.add_paragraph("Matching Weights: Semantic 55%, Domain 20%, Geo 15%, Boosts 10%")
    doc.add_paragraph("LLM Model: GPT-4o-mini (for reasoned analysis and CPV inference)")
    doc.add_paragraph("Embedding Model: text-embedding-3-small (1536 dimensions)")
    doc.add_paragraph("Database: Azure PostgreSQL with pgvector extension")

    report_path = "Grants_AI_Matching_Report.docx"
    doc.save(report_path)
    logger.info(f"Report generated: {report_path}")
    db.close()


if __name__ == "__main__":
    generate_report()
